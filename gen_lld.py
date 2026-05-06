from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin=Inches(1); s.bottom_margin=Inches(1)
    s.left_margin=Inches(1.2); s.right_margin=Inches(1.2)

def rf(run,sz=11,bold=False,color=None,italic=False):
    run.font.name='Calibri'; run.font.size=Pt(sz); run.font.bold=bold; run.font.italic=italic
    if color: run.font.color.rgb=RGBColor(*color)

def para(doc,txt='',sz=11,bold=False,color=None,align=None,sb=0,sa=4,italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if align: p.alignment=align
    if txt: r=p.add_run(txt); rf(r,sz,bold,color,italic)
    return p

def h1(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(txt); rf(r,15,True,(26,86,219))
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'8'); b.set(qn('w:space'),'4'); b.set(qn('w:color'),'1A56DB')
    pBdr.append(b); pPr.append(pBdr)

def h2(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(txt); rf(r,12,True,(15,23,42))

def h3(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(txt); rf(r,11,True,(30,64,175))

def bullet(doc,txt,indent=0.3):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Inches(indent)
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
    r=p.add_run(txt); rf(r,10.5,False,(51,65,85))

def numbered(doc,items,indent=0.3):
    for item in items:
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.left_indent=Inches(indent)
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        r=p.add_run(item); rf(r,10.5,False,(51,65,85))

def tbl(doc,headers,rows,col_widths=None):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=''; p=hdr[i].paragraphs[0]; r=p.add_run(h); rf(r,9,True,(255,255,255))
        tc=hdr[i]._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'1A56DB')
        tcPr.append(shd)
    for ri,row in enumerate(rows):
        cells=t.rows[ri+1].cells
        for ci,val in enumerate(row):
            cells[ci].text=''; p=cells[ci].paragraphs[0]; r=p.add_run(str(val)); rf(r,9.5)
            if ri%2==0:
                tc=cells[ci]._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
                shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'EFF6FF')
                tcPr.append(shd)
    if col_widths:
        for ri,row in enumerate(t.rows):
            for ci,cell in enumerate(row.cells): cell.width=Inches(col_widths[ci])
    doc.add_paragraph()

def callout(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(5)
    p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(txt); rf(r,10,False,(30,64,175),italic=True)

def divider(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(8)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'1'); b.set(qn('w:color'),'CBD5E1')
    pBdr.append(b); pPr.append(pBdr)

# TITLE PAGE
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80)
r=p.add_run('SmartSure Insurance Management System'); rf(r,22,True,(26,86,219))
p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=p2.add_run('Low-Level Design (LLD) Document'); rf(r2,16,False,(71,85,105))
p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_before=Pt(10)
r3=p3.add_run('Backend: ASP.NET Core 10 Microservices  |  Frontend: Angular 21'); rf(r3,11,False,(100,116,139))
p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
r4=p4.add_run('Version 1.0  |  May 2026  |  Complete and Authoritative'); rf(r4,10,False,(148,163,184))
doc.add_page_break()
print('title done')

# SECTION 1 - System Overview
h1(doc,"1. System Overview")
para(doc,"SmartSure is a full-stack insurance management platform composed of four ASP.NET Core 10 microservices behind an Ocelot API Gateway, with an Angular 21 SPA as the client. Each service owns its own SQL Server 2022 database. Asynchronous inter-service communication uses RabbitMQ for email notifications. Authentication is JWT HS256 validated independently in every service.",sz=11,color=(51,65,85))

h2(doc,"1.1 Architecture Summary")
tbl(doc,
    ["Component","Technology","Port","Responsibility"],
    [
        ["Angular SPA","Angular 21 / TypeScript","4200","Customer and Admin UI; standalone components, reactive forms, functional guards and interceptors"],
        ["API Gateway","Ocelot + SwaggerForOcelot","5000","Single entry-point: routing, CORS, JWT enforcement, aggregated Swagger UI"],
        ["Identity Service","ASP.NET Core 10","5265","Authentication, JWT (HS256), OTP email verification, BCrypt passwords, user management"],
        ["Policy Service","ASP.NET Core 10","5145","Policy types, premium calculation, policy creation, renewal, payment recording"],
        ["Claims Service","ASP.NET Core 10","5084","Claims lifecycle, document uploads, status state machine, claim statistics"],
        ["Admin Service","ASP.NET Core 10","5073","Dashboard aggregation, audit logs, report generation, user and claim management"],
        ["RabbitMQ","rabbitmq:3-management (Docker)","5672 / 15672","Async message bus for claim status email notifications"],
        ["SQL Server","mcr.microsoft.com/mssql/server:2022","1433","4 isolated databases, one per microservice, auto-migrated on startup"],
    ],
    col_widths=[1.4,1.8,0.7,2.9]
)

h2(doc,"1.2 Cross-Cutting Concerns")
bullet(doc,"Authentication: JWT Bearer (HS256). Token issued by Identity Service; all other services validate using shared SecretKey, Issuer, and Audience from appsettings / environment variables.")
bullet(doc,"Token Expiry: 8 hours. Configured via JwtSettings:ExpiryHours in appsettings.json.")
bullet(doc,"Logging: Serilog console logging enabled in all four services and the API Gateway.")
bullet(doc,"Error Handling: GlobalExceptionMiddleware in each service's API/Middlewares folder maps typed domain exceptions to structured JSON responses. Domain exceptions carry an int StatusCode property. The middleware logs domain exceptions as Warning and unexpected exceptions as Error.")
bullet(doc,"Repository Pattern: All services use repository interfaces (IAuthRepository, IPolicyRepository, IClaimRepository, IAdminRepository) with EF Core implementations. AdminService additionally uses IHttpClientFactory for downstream HTTP calls.")
bullet(doc,"Auto-Migration: All four services call db.Database.Migrate() on startup, creating and updating their databases automatically.")
bullet(doc,"Environment Config: Services read connection strings and JWT settings from environment variables injected by Docker Compose (ConnectionStrings__DefaultConnection, JwtSettings__SecretKey, etc.).")

h2(doc,"1.3 Clean Architecture Layers (per service)")
tbl(doc,
    ["Layer","Project Suffix","Responsibility"],
    [
        ["API","*.API","Controllers, middleware, Program.cs, Swagger configuration"],
        ["Application","*.Application","DTOs, service interfaces, service implementations, custom exceptions"],
        ["Domain","*.Domain","Entities, enums, repository interfaces"],
        ["Infrastructure","*.Infrastructure","EF Core DbContext, migrations, repository implementations, external services (email, RabbitMQ)"],
    ],
    col_widths=[1.2,1.5,4.0]
)

# SECTION 2 - Work Breakdown Structure (WBS)
h1(doc,"2. Work Breakdown Structure (WBS)")
para(doc,"The following WBS decomposes the SmartSure project into hierarchical deliverables and work packages, organized by major development phases.",sz=11,color=(51,65,85))

h2(doc,"2.1 Project Planning & Design")
bullet(doc,"1.1 Requirements Analysis: Gather functional and non-functional requirements from stakeholders")
bullet(doc,"1.2 System Architecture Design: Define microservices architecture, API Gateway, and communication patterns")
bullet(doc,"1.3 Database Schema Design: Design normalized schemas for Identity, Policy, Claims, and Admin databases")
bullet(doc,"1.4 Technology Stack Selection: Finalize ASP.NET Core 10, Angular 21, SQL Server 2022, RabbitMQ, Docker")
bullet(doc,"1.5 Security Design: Define JWT authentication, BCrypt password hashing, OTP verification flow")

h2(doc,"2.2 Backend Development - Identity Service")
bullet(doc,"2.1 User Entity & Repository: Implement User model, IAuthRepository interface, EF Core DbContext")
bullet(doc,"2.2 Registration & OTP Verification: Build send-otp, verify-register, resend-otp endpoints with 15-min expiry")
bullet(doc,"2.3 Login & JWT Generation: Implement BCrypt verification, JWT token generation with HS256")
bullet(doc,"2.4 Password Reset Flow: Build forgot-password/send-otp and forgot-password/reset endpoints")
bullet(doc,"2.5 Admin User Management: Implement get-all-users, get-user-by-id, update-user-status endpoints")
bullet(doc,"2.6 RabbitMQ Consumer: Build ClaimNotificationConsumer for claim status email notifications")
bullet(doc,"2.7 Email Service Integration: Integrate MailKit with Gmail SMTP for OTP and notification emails")

h2(doc,"2.3 Backend Development - Policy Service")
bullet(doc,"3.1 PolicyType Entity & CRUD: Implement PolicyType model with coverage details, admin CRUD endpoints")
bullet(doc,"3.2 Premium Calculation Engine: Build age factor (18-25, 26-40, 41-55, 56+) and duration factor logic")
bullet(doc,"3.3 Policy Creation: Implement POST /policies with Premium and Payment record creation")
bullet(doc,"3.4 Policy Renewal: Build renewal logic with 1-year extension, fixed durationFactor=0.10")
bullet(doc,"3.5 Payment Recording: Implement Payment entity with TransactionId generation (TXN-{ticks}, TXN-RENEW-{ms})")
bullet(doc,"3.6 My Policies Endpoint: Build GET /policies/my for customer policy listing")
bullet(doc,"3.7 Admin Policy Management: Implement policy status update and statistics endpoints")

h2(doc,"2.4 Backend Development - Claims Service")
bullet(doc,"4.1 Claim Entity & State Machine: Implement Claim model with status enum (Draft, Submitted, UnderReview, Approved, Rejected, Closed)")
bullet(doc,"4.2 Claim Creation & Submission: Build POST /claims (Draft) and POST /claims/{id}/submit (Submitted)")
bullet(doc,"4.3 Document Upload: Implement multipart/form-data upload to wwwroot/uploads/{claimId}/")
bullet(doc,"4.4 Document Deletion: Build DELETE /claims/{claimId}/documents/{docId} with Draft-only restriction")
bullet(doc,"4.5 Status Transition Validation: Enforce state machine rules (e.g., Draft → Submitted → UnderReview → Approved)")
bullet(doc,"4.6 My Claims Endpoint: Build GET /claims/my for customer claim listing with documents")
bullet(doc,"4.7 Admin Claims Statistics: Implement GET /claims/admin/stats with counts by status")

h2(doc,"2.5 Backend Development - Admin Service")
bullet(doc,"5.1 Dashboard Aggregation: Build parallel HTTP calls to Identity, Policy, Claims services for summary")
bullet(doc,"5.2 Claims Review: Implement GET /admin/claims and PUT /admin/claims/status with AdminLog creation")
bullet(doc,"5.3 User Management: Build GET /admin/users and PUT /admin/users/{id}/status endpoints")
bullet(doc,"5.4 Report Generation: Implement GET /admin/reports/generate with JSON serialization to AdminDB")
bullet(doc,"5.5 Audit Logging: Create AdminLog entity with action, targetType, targetId, notes, createdAt")
bullet(doc,"5.6 RabbitMQ Publisher: Build NotificationPublisher for fire-and-forget claim status notifications")
bullet(doc,"5.7 Background Task: Implement Task.Run with IHttpClientFactory for customer email fetch")

h2(doc,"2.6 API Gateway Configuration")
bullet(doc,"6.1 Ocelot Route Configuration: Define upstream/downstream routes for all four services")
bullet(doc,"6.2 JWT Validation Middleware: Configure shared JWT validation with SecretKey, Issuer, Audience")
bullet(doc,"6.3 CORS Policy: Set up AllowAngular policy for http://localhost:4200")
bullet(doc,"6.4 Swagger Aggregation: Integrate SwaggerForOcelot to aggregate all service Swagger UIs")
bullet(doc,"6.5 Environment-Specific Config: Create ocelot.json (Development) and ocelot.Docker.json (Docker)")

h2(doc,"2.7 Frontend Development - Customer Portal")
bullet(doc,"7.1 Authentication Module: Build login, register, OTP verification, forgot-password components")
bullet(doc,"7.2 Policy Purchase Flow: Create 3-step wizard (select type, calculate premium, Razorpay payment)")
bullet(doc,"7.3 My Policies Page: Build policy listing with status badges, renewal button, payment history")
bullet(doc,"7.4 Claims Management: Create claim form, document upload, submission, and status tracking")
bullet(doc,"7.5 Dashboard: Build customer dashboard with active policies, pending claims, recent payments")
bullet(doc,"7.6 Auth Interceptor: Implement JWT token attachment to all HTTP requests")
bullet(doc,"7.7 Route Guards: Build canActivate guards for CUSTOMER and ADMIN role-based routing")

h2(doc,"2.8 Frontend Development - Admin Portal")
bullet(doc,"8.1 Admin Dashboard: Build aggregated statistics cards (users, policies, claims, revenue)")
bullet(doc,"8.2 Claims Review Panel: Create slide-in panel with claim details, documents, status update form")
bullet(doc,"8.3 User Management: Build user listing table with activate/deactivate toggle")
bullet(doc,"8.4 Policy Type Management: Create CRUD interface for PolicyType with coverage details")
bullet(doc,"8.5 Reports Generation: Build report type selector and download functionality")
bullet(doc,"8.6 Audit Logs Viewer: Create AdminLog listing with filters by action, targetType, date range")

h2(doc,"2.9 Integration & Messaging")
bullet(doc,"9.1 RabbitMQ Setup: Configure rabbitmq:3-management Docker container with durable queues")
bullet(doc,"9.2 Claim Status Notification Flow: Integrate AdminService publisher → IdentityService consumer")
bullet(doc,"9.3 Email Template Design: Create HTML email templates with color-coded status badges")
bullet(doc,"9.4 Connection Resilience: Implement auto-reconnect logic with 5-second retry delay")
bullet(doc,"9.5 Message Persistence: Configure persistent messages and durable queues for reliability")

h2(doc,"2.10 Testing")
bullet(doc,"10.1 Identity Service Tests: Write 31 NUnit tests covering OTP flow, login, password reset (98% coverage)")
bullet(doc,"10.2 Policy Service Tests: Write 34 NUnit tests covering premium calculation, renewal, payments (97% coverage)")
bullet(doc,"10.3 Claims Service Tests: Write 30 NUnit tests covering state machine, documents, access control (95% coverage)")
bullet(doc,"10.4 Admin Service Tests: Write 22 NUnit tests covering aggregation, status updates, audit logs (90% coverage)")
bullet(doc,"10.5 Integration Testing: Test end-to-end flows (registration → policy purchase → claim submission → admin review)")
bullet(doc,"10.6 API Gateway Testing: Verify routing, JWT validation, CORS, Swagger aggregation")

h2(doc,"2.11 Infrastructure & Deployment")
bullet(doc,"11.1 Docker Compose Configuration: Define 8 services (SQL Server, RabbitMQ, 4 microservices, Gateway, Frontend)")
bullet(doc,"11.2 Database Migration: Implement auto-migration on startup (db.Database.Migrate()) for all services")
bullet(doc,"11.3 Environment Variables: Configure connection strings, JWT settings, RabbitMQ credentials via docker-compose.yml")
bullet(doc,"11.4 Health Checks: Add SQL Server and RabbitMQ health checks with depends_on conditions")
bullet(doc,"11.5 Nginx Configuration: Set up Nginx reverse proxy for Angular SPA with /gateway/ API routing")
bullet(doc,"11.6 Production Deployment: Deploy to cloud (Azure/AWS) with managed SQL, RabbitMQ, container orchestration")

h2(doc,"2.12 Documentation")
bullet(doc,"12.1 Low-Level Design (LLD): Create comprehensive LLD document with architecture, APIs, data models, flows")
bullet(doc,"12.2 API Documentation: Generate Swagger/OpenAPI specs for all four microservices")
bullet(doc,"12.3 Sequence Diagrams: Create diagrams for key flows (registration, policy purchase, claim review, renewal)")
bullet(doc,"12.4 ER Diagram: Design entity-relationship diagram showing all 10 entities across 4 databases")
bullet(doc,"12.5 Deployment Guide: Write Docker setup instructions, environment configuration, troubleshooting")
bullet(doc,"12.6 User Manual: Create end-user documentation for customer and admin portal features")

divider(doc)

# Then renumber Section 2 to Section 3
h1(doc,"3. API Gateway (Ocelot)")  # was Section 2


para(doc,"Port: 5000  |  Technology: Ocelot 24.1.0 + SwaggerForOcelot",sz=11,bold=True,color=(15,23,42))
para(doc,"Single entry-point for the Angular SPA. Handles routing, CORS, JWT validation, and aggregated Swagger UI. Loads ocelot.json in Development and ocelot.Docker.json in Docker environments.",sz=11,color=(51,65,85))

h2(doc,"3.1 Route Table")
tbl(doc,
    ["Upstream Path (Client calls)","Downstream Path","Service Port","Auth Required"],
    [
        ["POST /gateway/auth/register","/api/auth/register","Identity 5265","No"],
        ["POST /gateway/auth/login","/api/auth/login","Identity 5265","No"],
        ["POST /gateway/auth/send-otp","/api/auth/send-otp","Identity 5265","No"],
        ["POST /gateway/auth/verify-register","/api/auth/verify-register","Identity 5265","No"],
        ["POST /gateway/auth/resend-otp","/api/auth/resend-otp","Identity 5265","No"],
        ["POST /gateway/auth/forgot-password/send-otp","/api/auth/forgot-password/send-otp","Identity 5265","No"],
        ["POST /gateway/auth/forgot-password/reset","/api/auth/forgot-password/reset","Identity 5265","No"],
        ["GET|POST|PUT|DELETE /gateway/auth/{everything}","/api/auth/{everything}","Identity 5265","Bearer JWT"],
        ["POST /gateway/policies/{id}/renew","/api/policies/{id}/renew","Policy 5145","Bearer JWT"],
        ["GET|POST|PUT|DELETE /gateway/policies/{everything}","/api/policies/{everything}","Policy 5145","Bearer JWT"],
        ["GET|POST|PUT|DELETE /gateway/claims/{everything}","/api/claims/{everything}","Claims 5084","Bearer JWT"],
        ["GET|POST|PUT|DELETE /gateway/admin/{everything}","/api/admin/{everything}","Admin 5073","Bearer JWT"],
    ],
    col_widths=[2.5,2.0,1.2,1.0]
)

h2(doc,"3.2 CORS Policy")
para(doc,"Policy name: AllowAngular. Permits all HTTP methods and headers from http://localhost:4200 (local development). In Docker, the Angular frontend is served by Nginx which proxies /gateway/ calls to the api-gateway container on the internal Docker network.",sz=11,color=(51,65,85))

h2(doc,"3.3 Swagger Aggregation")
para(doc,"SwaggerForOcelot aggregates Swagger JSON from all four downstream services and presents a unified API explorer at /swagger on the gateway. Swagger endpoints:",sz=11,color=(51,65,85))
bullet(doc,"Identity Service: http://localhost:5265/swagger/v1/swagger.json")
bullet(doc,"Policy Service: http://localhost:5145/swagger/v1/swagger.json")
bullet(doc,"Claims Service: http://localhost:5084/swagger/v1/swagger.json")
bullet(doc,"Admin Service: http://localhost:5073/swagger/v1/swagger.json")

print("section 1-2 done")

# SECTION 3 - Identity Service
h1(doc,"4. Identity Service")
para(doc,"Port: 5265  |  Database: SmartSure_IdentityDB  |  Route prefix: /api/auth",sz=11,bold=True,color=(15,23,42))

h2(doc,"4.1 Data Models")
h3(doc,"User (Primary Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["FullName","string","Required; stored in Title Case"],
        ["Email","string","Required; stored lowercase; unique"],
        ["PasswordHash","string","BCrypt hash of the user password"],
        ["Role","string","CUSTOMER or ADMIN"],
        ["IsActive","bool","Default true; set false by admin to deactivate account"],
        ["CreatedAt","DateTime (UTC)","Auto-set at account creation"],
    ],
    col_widths=[1.5,1.2,4.0]
)
h3(doc,"OtpVerification (Supporting Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["Email","string","Email address the OTP was sent to"],
        ["OtpCode","string","6-digit numeric code"],
        ["ExpiresAt","DateTime (UTC)","15 minutes from creation"],
        ["IsUsed","bool","Set true after successful verification; prevents reuse"],
    ],
    col_widths=[1.5,1.2,4.0]
)

h2(doc,"4.2 API Endpoints")
tbl(doc,
    ["Verb","Path","Auth","Description"],
    [
        ["POST","/api/auth/register","None","Triggers OTP send; returns 202 Accepted with requiresOtpVerification=true"],
        ["POST","/api/auth/send-otp","None","Sends 6-digit OTP to email; upserts OtpVerification record with 15-min expiry"],
        ["POST","/api/auth/verify-register","None","Validates OTP; creates User account; returns JWT token"],
        ["POST","/api/auth/resend-otp","None","Regenerates and resends OTP for the given email"],
        ["POST","/api/auth/login","None","Validates BCrypt hash; checks IsActive; returns JWT token"],
        ["POST","/api/auth/forgot-password/send-otp","None","Sends password reset OTP to registered email"],
        ["POST","/api/auth/forgot-password/reset","None","Validates OTP; updates BCrypt password hash"],
        ["GET","/api/auth/profile","[Authorize]","Returns authenticated user profile from JWT NameIdentifier claim"],
        ["GET","/api/auth/admin/users","[ADMIN]","Returns all users with Id, FullName, Email, Role, IsActive, CreatedAt"],
        ["GET","/api/auth/admin/users/{userId}","[ADMIN]","Returns profile for a specific user by ID"],
        ["GET","/api/auth/admin/users/count","[ADMIN]","Returns totalUsers, activeUsers, inactiveUsers counts"],
        ["PUT","/api/auth/admin/users/{userId}/status","[ADMIN]","Sets IsActive flag for a user (activate or deactivate)"],
    ],
    col_widths=[0.6,2.8,0.9,2.5]
)

h2(doc,"4.3 Service Layer — AuthService")
bullet(doc,"SendRegistrationOtpAsync: Checks email uniqueness (throws EmailAlreadyRegisteredException if exists). Generates 6-digit OTP via RandomNumberGenerator.GetInt32. Calls IAuthRepository.UpsertOtpAsync to store with 15-minute expiry. Calls IEmailService.SendOtpEmailAsync.")
bullet(doc,"VerifyRegistrationOtpAsync: Checks email uniqueness again. Retrieves latest OTP record. Throws OtpNotFoundException if null or IsUsed=true. Throws OtpExpiredException if ExpiresAt < UtcNow. Throws InvalidOtpException if code mismatch (ordinal comparison). Creates User with BCrypt-hashed password and role CUSTOMER. Marks OTP as used.")
bullet(doc,"LoginAsync: Retrieves user by email. Throws InvalidCredentialsException if null or BCrypt.Verify fails. Throws AccountDeactivatedException if IsActive=false. Returns AuthResponseDto (token generated by JwtHelper in controller).")
bullet(doc,"SendPasswordResetOtpAsync: Looks up user by email. Throws UserNotFoundException if not found. Generates OTP, upserts record, sends email.")
bullet(doc,"ResetPasswordAsync: Validates OTP (same checks as registration). Hashes new password with BCrypt. Calls UpdatePasswordAsync. Marks OTP as used.")
bullet(doc,"GetAllUsersAsync / GetProfileAsync / UpdateUserStatusAsync: Admin operations delegating to IAuthRepository.")

h2(doc,"4.4 JWT Token Generation (JwtHelper)")
bullet(doc,"Algorithm: HS256 (HMAC-SHA256)")
bullet(doc,"Claims: ClaimTypes.NameIdentifier (user Id as string), ClaimTypes.Email, ClaimTypes.Role")
bullet(doc,"Expiry: Configured via JwtSettings:ExpiryHours (default 8 hours)")
bullet(doc,"Signing key: UTF-8 bytes of JwtSettings:SecretKey")
bullet(doc,"Issuer / Audience: JwtSettings:Issuer and JwtSettings:Audience (shared across all services)")

h2(doc,"4.5 RabbitMQ Consumer — ClaimNotificationConsumer")
bullet(doc,"Type: BackgroundService (IHostedService)")
bullet(doc,"Queue: claim.status.notification (durable, persistent messages)")
bullet(doc,"On message received: Deserialises ClaimStatusNotificationDto. Resolves scoped IEmailService via IServiceScopeFactory. Calls SendClaimStatusEmailAsync with colour-coded HTML template.")
bullet(doc,"Reconnection: Auto-reconnects on RabbitMQ disconnect with 5-second retry delay.")
bullet(doc,"Email colours: Approved=green, Rejected=red, UnderReview=purple, Closed=grey")

h2(doc,"4.6 Custom Exceptions")
tbl(doc,
    ["Exception Class","HTTP Status","Thrown When"],
    [
        ["EmailAlreadyRegisteredException","409 Conflict","Registration attempted with an already-registered email"],
        ["UserNotFoundException","404 Not Found","Password reset OTP requested for non-existent email"],
        ["InvalidCredentialsException","401 Unauthorized","Login with wrong email or wrong password"],
        ["AccountDeactivatedException","401 Unauthorized","Login attempted on a deactivated account"],
        ["OtpNotFoundException","400 Bad Request","OTP verification attempted with no active OTP record"],
        ["OtpExpiredException","400 Bad Request","OTP submitted after the 15-minute expiry window"],
        ["InvalidOtpException","400 Bad Request","OTP code submitted does not match the stored code"],
    ],
    col_widths=[2.4,1.3,3.1]
)

print("section 4 done")

# SECTION 4 - Policy Service
h1(doc,"5. Policy Service")
para(doc,"Port: 5145  |  Database: SmartSure_PolicyDB  |  Route prefix: /api/policies",sz=11,bold=True,color=(15,23,42))

h2(doc,"5.1 Data Models")
h3(doc,"PolicyType (Configuration Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["Name","string","e.g. Health Insurance, Auto Insurance"],
        ["Description","string","Short description shown to customers"],
        ["BaseAmount","decimal","Base premium amount before age/duration factors"],
        ["IsActive","bool","Only active types shown on buy-policy page"],
        ["CoverageDetails","string","What is covered (shown in buy-policy UI)"],
        ["Exclusions","string","What is not covered (shown in buy-policy UI)"],
        ["ClaimLimit","decimal","Maximum claim payout amount"],
        ["MinAge / MaxAge","int","Eligible age range for this policy type"],
        ["DurationMonths","int","Standard policy duration in months"],
        ["RiskCategory","string","Low, Medium, or High"],
        ["AutoRenewal","bool","Whether the policy supports auto-renewal"],
        ["GracePeriodDays","int","Grace period after expiry before cancellation"],
    ],
    col_widths=[1.8,1.0,3.9]
)
h3(doc,"Policy (Core Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["UserId","int","References Identity user (no cross-service FK constraint)"],
        ["PolicyTypeId","int (FK)","References PolicyType"],
        ["PolicyNumber","string","Unique identifier: POL-{DateTime.UtcNow.Ticks}"],
        ["Status","PolicyStatus enum","Active, Expired, Cancelled, Draft"],
        ["StartDate / EndDate","DateTime (UTC)","Coverage period"],
        ["PremiumAmount","decimal","Final calculated premium"],
        ["IsRenewed","bool","True if this policy has been renewed at least once"],
        ["RenewalCount","int","Number of times this policy has been renewed"],
        ["CreatedAt","DateTime (UTC)","Auto-set at creation"],
        ["Premium","Navigation","One-to-one: stores BaseAmount, AgeFactor, DurationFactor, FinalAmount"],
    ],
    col_widths=[1.8,1.0,3.9]
)
h3(doc,"Payment (Supporting Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["PolicyId","int (FK)","References Policy"],
        ["UserId","int","References Identity user"],
        ["Amount","decimal","Payment amount"],
        ["PaymentMethod","string","Online (Razorpay)"],
        ["Status","string","Success, Failed, Pending"],
        ["TransactionId","string","TXN-{ticks} for new policies; TXN-RENEW-{ms} for renewals"],
        ["PaymentDate","DateTime (UTC)","Auto-set at creation"],
    ],
    col_widths=[1.8,1.0,3.9]
)

h2(doc,"5.2 API Endpoints")
tbl(doc,
    ["Verb","Path","Auth","Description"],
    [
        ["GET","/api/policies/types","None (public)","List all active policy types with full coverage details"],
        ["GET","/api/policies/types/{id}","None (public)","Get a specific policy type by ID"],
        ["POST","/api/policies/calculate-premium","[Authorize]","Calculate premium breakdown; no database write"],
        ["POST","/api/policies","[Authorize]","Create policy with Active status; record payment"],
        ["GET","/api/policies/my","[Authorize]","Get all policies for the authenticated user"],
        ["GET","/api/policies/{id}","[Authorize]","Get a specific policy by ID"],
        ["GET","/api/policies/{id}/payment","[Authorize]","Get payment record for a policy"],
        ["GET","/api/policies/my/payments","[CUSTOMER]","Get all payment records for the authenticated user"],
        ["POST","/api/policies/{id}/renew","[Authorize]","Renew an Active or Expired policy for 1 year"],
        ["PUT","/api/policies/{id}/status","[ADMIN]","Update policy status (Admin only)"],
        ["GET","/api/policies/admin/count","[ADMIN]","Returns totalPolicies and totalRevenue"],
        ["GET","/api/policies/admin/types","[ADMIN]","List all policy types including inactive ones with enrolled count"],
        ["POST","/api/policies/admin/types","[ADMIN]","Create a new policy type"],
        ["PUT","/api/policies/admin/types/{id}","[ADMIN]","Update an existing policy type"],
        ["PUT","/api/policies/admin/types/{id}/toggle","[ADMIN]","Toggle IsActive status of a policy type"],
        ["DELETE","/api/policies/admin/types/{id}","[ADMIN]","Delete a policy type"],
        ["GET","/api/policies/admin/types/{id}/stats","[ADMIN]","Get total policies, active policies, and total premium for a type"],
    ],
    col_widths=[0.6,2.8,0.9,2.5]
)

h2(doc,"5.3 Premium Calculation Logic")
para(doc,"Implemented in PolicyAppService.CalculatePremiumAsync. The formula is additive:",sz=11,color=(51,65,85))
callout(doc,"Final Premium = Base Amount + Age Factor Amount + Duration Factor Amount")
tbl(doc,
    ["Parameter","Rule"],
    [
        ["Age Factor — 18 to 25 years","+10% of Base Amount"],
        ["Age Factor — 26 to 40 years","+0% (no surcharge)"],
        ["Age Factor — 41 to 55 years","+20% of Base Amount"],
        ["Age Factor — 56 years and above","+50% of Base Amount"],
        ["Duration Factor — 1 Year","+(1-1)+0.10 = 0.10 = +10% of Base Amount"],
        ["Duration Factor — 2 Years","+(2-1)+0.10 = 1.10 = +110% of Base Amount"],
        ["Duration Factor — 3 Years","+(3-1)+0.10 = 2.10 = +210% of Base Amount"],
        ["Duration Factor — N Years","+(N-1)+0.10 = +(N-0.9)*100% of Base Amount"],
        ["Duration Calculation","Math.Ceiling(totalDays / 365.0) — partial years round up"],
        ["Renewal Duration Factor","Always fixed at 0.10 (1-year rate) regardless of original duration"],
    ],
    col_widths=[2.8,3.9]
)

h2(doc,"5.4 Service Layer — PolicyAppService")
bullet(doc,"GetAllPolicyTypesAsync: Returns all active policy types with full coverage details mapped to PolicyTypeResponseDto.")
bullet(doc,"CalculatePremiumAsync: Looks up PolicyType. Calculates years = Ceiling(days/365). Applies age and duration factors. Returns PremiumResponseDto with BaseAmount, AgeFactor, AgeFactorAmount, DurationFactor, DurationFactorAmount, DurationYears, FinalAmount, AgeGroup, FormulaExplanation.")
bullet(doc,"CreatePolicyAsync: Calls CalculatePremiumAsync. Creates Policy with status Active and PolicyNumber=POL-{ticks}. Creates Premium record. Creates Payment record with TransactionId=TXN-{ticks}.")
bullet(doc,"RenewPolicyAsync: Validates ownership (throws PolicyAccessDeniedException if UserId mismatch). Validates status is Active or Expired (throws PolicyNotRenewableException otherwise). Calculates renewal premium with fixed durationFactor=0.10. Updates Policy dates, status, IsRenewed=true, RenewalCount+=1. Creates Payment with TransactionId=TXN-RENEW-{ms}.")
bullet(doc,"UpdatePolicyStatusAsync: Parses status string to PolicyStatus enum (throws InvalidPolicyStatusException if invalid). Updates policy status.")

h2(doc,"5.5 Custom Exceptions")
tbl(doc,
    ["Exception Class","HTTP Status","Thrown When"],
    [
        ["PolicyTypeNotFoundException","404 Not Found","Policy type ID not found in database"],
        ["PolicyNotFoundException","404 Not Found","Policy ID not found in database"],
        ["PaymentNotFoundException","404 Not Found","No payment record found for the given policy ID"],
        ["PolicyAccessDeniedException","403 Forbidden","Renewal attempted on a policy owned by a different user"],
        ["PolicyNotRenewableException","400 Bad Request","Renewal attempted on a policy in Draft or Cancelled status"],
        ["InvalidPolicyStatusException","400 Bad Request","Admin submits an unrecognised status string"],
    ],
    col_widths=[2.4,1.3,3.1]
)

print("section 5 done")

# SECTION 5 - Claims Service
h1(doc,"6. Claims Service")
para(doc,"Port: 5084  |  Database: SmartSure_ClaimsDB  |  Route prefix: /api/claims",sz=11,bold=True,color=(15,23,42))

h2(doc,"6.1 Data Models")
h3(doc,"Claim (Core Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["PolicyId","int","References Policy Service (no cross-service FK constraint)"],
        ["CustomerId","int","References Identity user"],
        ["ClaimNumber","string","Unique identifier: CLM-{DateTime.UtcNow.Ticks}"],
        ["IncidentDate","DateTime","Date the incident occurred"],
        ["Description","string","Required; trimmed on creation"],
        ["Status","ClaimStatus enum","Draft, Submitted, UnderReview, Approved, Rejected, Closed"],
        ["AdminNote","string?","Optional note added by admin on status update"],
        ["CreatedAt / UpdatedAt","DateTime (UTC)","Audit timestamps; UpdatedAt refreshed on every change"],
        ["ClaimDocuments","Navigation","Collection of ClaimDocument records"],
    ],
    col_widths=[1.8,1.0,3.9]
)
h3(doc,"ClaimDocument (Supporting Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["ClaimId","int (FK)","References Claim"],
        ["FileName","string","Original file name"],
        ["FilePath","string","Server path under wwwroot/uploads/{claimId}/"],
        ["FileType","string","MIME type (e.g. application/pdf, image/jpeg)"],
        ["FileSize","long","File size in bytes"],
        ["UploadedAt","DateTime (UTC)","Auto-set at upload"],
    ],
    col_widths=[1.8,1.0,3.9]
)

h2(doc,"6.2 Claim Status Lifecycle")
tbl(doc,
    ["Status","Triggered By","Allowed Next Statuses","Notes"],
    [
        ["Draft","Customer creates claim","Submitted","Starting state; documents can be uploaded and deleted"],
        ["Submitted","Customer submits claim","UnderReview","No edits allowed; RabbitMQ event published"],
        ["UnderReview","Admin sets review","Approved, Rejected","Admin is actively reviewing the claim"],
        ["Approved","Admin approves","Closed","Admin note recorded; email notification sent to customer"],
        ["Rejected","Admin rejects","Closed","Admin note recorded; email notification sent to customer"],
        ["Closed","Admin closes","None (terminal)","Final state; no further transitions allowed"],
    ],
    col_widths=[1.1,1.5,1.6,2.6]
)

h2(doc,"6.3 API Endpoints")
tbl(doc,
    ["Verb","Path","Auth","Description"],
    [
        ["POST","/api/claims","[CUSTOMER]","Create claim in Draft status"],
        ["POST","/api/claims/{id}/submit","[CUSTOMER]","Transition Draft to Submitted; publish RabbitMQ event"],
        ["POST","/api/claims/{id}/documents","[CUSTOMER]","Upload document (multipart/form-data); save to wwwroot/uploads/{id}/"],
        ["DELETE","/api/claims/{claimId}/documents/{docId}","[CUSTOMER]","Delete document (only allowed on Draft claims)"],
        ["GET","/api/claims/my","[CUSTOMER]","Get all claims for the authenticated customer"],
        ["GET","/api/claims/{id}","[Authorize]","Get a specific claim with documents"],
        ["GET","/api/claims","[ADMIN]","Get all claims across all customers"],
        ["PUT","/api/claims/{id}/status","[ADMIN]","Update claim status with admin note; enforces state machine"],
        ["GET","/api/claims/admin/stats","[ADMIN]","Returns counts by status: total, draft, submitted, underReview, approved, rejected, closed"],
    ],
    col_widths=[0.6,2.8,0.9,2.5]
)

h2(doc,"6.4 Service Layer — ClaimAppService")
bullet(doc,"CreateClaimAsync: Creates Claim with status Draft, ClaimNumber=CLM-{ticks}, trimmed Description.")
bullet(doc,"SubmitClaimAsync: Validates claim exists (ClaimNotFoundException). Validates ownership (ClaimAccessDeniedException). Validates status is Draft (ClaimAlreadySubmittedException). Transitions to Submitted.")
bullet(doc,"UpdateClaimStatusAsync: Validates claim exists. Parses status string to ClaimStatus enum (InvalidClaimStatusException if invalid). Enforces state machine via switch expression (InvalidClaimStatusTransitionException if invalid transition). Updates status and AdminNote.")
bullet(doc,"AddDocumentAsync: Validates claim exists. Saves file via FileStorageService to wwwroot/uploads/{claimId}/. Creates ClaimDocument record. Returns FileUrl with wwwroot prefix stripped.")
bullet(doc,"DeleteDocumentAsync: Validates claim exists and ownership. Validates status is Draft (ClaimNotEditableException). Validates document exists (ClaimDocumentNotFoundException). Validates document belongs to claim (DocumentClaimMismatchException). Deletes physical file if it exists. Deletes ClaimDocument record.")
bullet(doc,"GetClaimsStatsAsync: Returns anonymous object with counts for each ClaimStatus value.")

h2(doc,"6.5 Custom Exceptions")
tbl(doc,
    ["Exception Class","HTTP Status","Thrown When"],
    [
        ["ClaimNotFoundException","404 Not Found","Claim ID not found in database"],
        ["ClaimDocumentNotFoundException","404 Not Found","Document ID not found in database"],
        ["ClaimAccessDeniedException","403 Forbidden","Customer attempts to act on a claim they do not own"],
        ["ClaimAlreadySubmittedException","400 Bad Request","Submit attempted on a claim not in Draft status"],
        ["InvalidClaimStatusException","400 Bad Request","Admin submits an unrecognised status string"],
        ["InvalidClaimStatusTransitionException","400 Bad Request","Status transition violates the state machine rules"],
        ["ClaimNotEditableException","400 Bad Request","Document operation attempted on a non-Draft claim"],
        ["DocumentClaimMismatchException","400 Bad Request","Document ID does not belong to the specified claim"],
    ],
    col_widths=[2.4,1.3,3.1]
)

print("section 6 done")

# SECTION 6 - Admin Service
h1(doc,"7. Admin Service")
para(doc,"Port: 5073  |  Database: SmartSure_AdminDB  |  Route prefix: /api/admin",sz=11,bold=True,color=(15,23,42))
para(doc,"The Admin Service is an orchestrator, not a data owner. It aggregates data from the three domain services via direct HTTP calls (forwarding the JWT Bearer token) and maintains its own AdminDB for audit logs and reports.",sz=11,color=(51,65,85))

h2(doc,"7.1 Data Models")
h3(doc,"AdminLog (Audit Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["AdminId","int","ID of the admin who performed the action"],
        ["Action","string","UpdateClaimStatus, UpdateUserStatus, GenerateReport"],
        ["TargetType","string","Claim, User, Report"],
        ["TargetId","int","ID of the entity acted upon"],
        ["Notes","string?","Optional notes (e.g. admin note on claim, new IsActive value)"],
        ["CreatedAt","DateTime (UTC)","Timestamp of the action"],
    ],
    col_widths=[1.5,1.0,4.2]
)
h3(doc,"Report (Reporting Entity)")
tbl(doc,
    ["Property","Type","Notes"],
    [
        ["Id","int (PK)","Auto-increment primary key"],
        ["ReportType","ReportType enum","ClaimsSummary, UserSummary, PolicySummary, RevenueSummary"],
        ["GeneratedBy","int","Admin user ID"],
        ["GeneratedAt","DateTime (UTC)","Timestamp of generation"],
        ["Data","string","Serialised JSON of the DashboardSummaryDto at time of generation"],
    ],
    col_widths=[1.5,1.0,4.2]
)

h2(doc,"7.2 API Endpoints")
tbl(doc,
    ["Verb","Path","Auth","Description"],
    [
        ["GET","/api/admin/dashboard","[ADMIN]","Aggregates totalUsers, activeUsers, totalPolicies, totalRevenue, totalClaims, pendingClaims, approvedClaims, rejectedClaims, closedClaims"],
        ["GET","/api/admin/claims","[ADMIN]","Proxies GET /api/claims to ClaimsService; returns all claims"],
        ["GET","/api/admin/claims/pending","[ADMIN]","Returns claims with status Submitted or UnderReview"],
        ["PUT","/api/admin/claims/status","[ADMIN]","Updates claim status in ClaimsService; logs action; fires RabbitMQ notification"],
        ["GET","/api/admin/users","[ADMIN]","Proxies GET /api/auth/admin/users to IdentityService"],
        ["PUT","/api/admin/users/{userId}/status","[ADMIN]","Proxies PUT to IdentityService; logs action in AdminLog"],
        ["GET","/api/admin/reports/generate","[ADMIN]","Generates and persists a report; logs action"],
        ["GET","/api/admin/logs","[ADMIN]","Returns all AdminLog records"],
    ],
    col_widths=[0.6,2.8,0.9,2.5]
)

h2(doc,"7.3 Service Layer — AdminAppService")
bullet(doc,"GetDashboardSummaryAsync: Fires three parallel HTTP calls using Task.WhenAll. Calls IdentityService /api/auth/admin/users/count, PolicyService /api/policies/admin/count, ClaimsService /api/claims/admin/stats. Aggregates into DashboardSummaryDto. If any downstream call fails, logs a Warning and returns partial data.")
bullet(doc,"UpdateClaimStatusAsync: Pre-fetches claim details (claimNumber, customerId, oldStatus) from ClaimsService. Sends PUT to ClaimsService /api/claims/{id}/status. Creates AdminLog entry. Fires fire-and-forget background task: fetches customer email from IdentityService using IHttpClientFactory (scope-independent), publishes ClaimStatusNotificationDto to RabbitMQ queue claim.status.notification.")
bullet(doc,"UpdateUserStatusAsync: Sends PUT to IdentityService /api/auth/admin/users/{userId}/status. Creates AdminLog entry.")
bullet(doc,"GenerateReportAsync: Calls GetDashboardSummaryAsync. Parses reportType to ReportType enum (defaults to ClaimsSummary if unrecognised). Creates Report record with serialised JSON data. Creates AdminLog entry.")
bullet(doc,"GetAdminLogsAsync: Returns all AdminLog records from AdminDB.")

h2(doc,"7.4 Fire-and-Forget Notification Pattern")
para(doc,"The RabbitMQ notification is published as a fire-and-forget background task to ensure the API response is never blocked by email delivery:",sz=11,color=(51,65,85))
numbered(doc,[
    "Admin calls PUT /api/admin/claims/status.",
    "AdminService pre-fetches claim details from ClaimsService.",
    "AdminService sends the status update to ClaimsService.",
    "AdminService creates an AdminLog entry.",
    "AdminService captures the JWT token from the current HttpContext (before the scope is disposed).",
    "AdminService starts a background Task using IHttpClientFactory.CreateClient() with its own 25-second timeout and a separate CancellationTokenSource(30s).",
    "Background task fetches customer email from IdentityService.",
    "Background task publishes ClaimStatusNotificationDto to the claim.status.notification RabbitMQ queue.",
    "IdentityService ClaimNotificationConsumer receives the message and sends the HTML email.",
    "If RabbitMQ is unavailable, the failure is logged as a Warning and the API response is unaffected.",
])

h2(doc,"7.5 Custom Exceptions")
tbl(doc,
    ["Exception Class","HTTP Status","Thrown When"],
    [
        ["AdminUserNotFoundException","404 Not Found","User ID cannot be resolved from IdentityService"],
        ["AdminClaimNotFoundException","404 Not Found","Claim ID cannot be resolved from ClaimsService"],
        ["DownstreamServiceException","502 Bad Gateway","Required downstream HTTP call fails"],
        ["InvalidReportTypeException","400 Bad Request","Unrecognised report type string submitted"],
    ],
    col_widths=[2.4,1.3,3.1]
)

print("section 7 done")

# SECTION 7 - Frontend
h1(doc,"8. Frontend (Angular 21 SPA)")
para(doc,"Port: 4200 (local) / 80 via Nginx (Docker)  |  Technology: Angular 21 / TypeScript",sz=11,bold=True,color=(15,23,42))
para(doc,"A standalone Angular 21 application. All components use standalone: true. Angular Router handles client-side navigation. Functional HTTP interceptors inject JWT. Functional route guards enforce access control. All API calls target the gateway at http://localhost:5000/gateway.",sz=11,color=(51,65,85))

h2(doc,"8.1 Route Table")
tbl(doc,
    ["Route Path","Component","Guard","Notes"],
    [
        ["/","LandingComponent","None","Public landing page with insurance showcase and CTAs"],
        ["/home","LandingComponent","None","Alias for landing page"],
        ["/auth/login","LoginComponent","None","Email/password login with animated left panel"],
        ["/auth/register","RegisterComponent","None","Two-step OTP registration flow"],
        ["/auth/forgot-password","ForgotPasswordComponent","None","Three-step password reset flow"],
        ["/customer","CustomerLayoutComponent","authGuard","Parent route; requires valid JWT"],
        ["/customer/dashboard","CustomerDashboardComponent","authGuard","Customer overview"],
        ["/customer/policies","PolicyListComponent","authGuard","List of customer policies"],
        ["/customer/policies/:id","PolicyDetailComponent","authGuard","Individual policy details"],
        ["/customer/claims","ClaimListComponent","authGuard","List of customer claims"],
        ["/customer/claims/:id","ClaimDetailComponent","authGuard","Individual claim details with documents"],
        ["/customer/buy-policy","BuyPolicyComponent","authGuard","3-step wizard: details, premium, confirm+pay"],
        ["/customer/initiate-claim","InitiateClaimComponent","authGuard","Create new claim form"],
        ["/admin","AdminLayoutComponent","adminGuard","Parent route; requires ADMIN role"],
        ["/admin/dashboard","AdminDashboardComponent","adminGuard","Aggregated KPI dashboard"],
        ["/admin/claims","AdminClaimsComponent","adminGuard","All claims with status update panel"],
        ["/admin/users","AdminUsersComponent","adminGuard","All users with activate/deactivate toggle"],
        ["/admin/reports","AdminReportsComponent","adminGuard","Report generation and audit logs"],
        ["/admin/policies","AdminPolicyManagementComponent","adminGuard","Policy type CRUD management"],
    ],
    col_widths=[2.0,1.8,0.9,2.1]
)

h2(doc,"8.2 Route Guards")
bullet(doc,"authGuard: Checks if JWT token exists in localStorage via TokenService.isLoggedIn(). If missing, redirects to /auth/login.")
bullet(doc,"adminGuard: Checks token exists, then validates role is ADMIN via TokenService.isAdmin(). No token redirects to /auth/login. Non-admin role redirects to /customer/dashboard.")

h2(doc,"8.3 HTTP Interceptor — AuthInterceptor")
bullet(doc,"Attaches JWT Bearer token to all outgoing HTTP requests.")
bullet(doc,"Injects TokenService.getToken(). If a token exists, clones the request and adds Authorization: Bearer {token} header.")
bullet(doc,"Applied globally via provideHttpClient(withInterceptors([authInterceptor])) in app.config.ts.")

h2(doc,"8.4 Core Services")
h3(doc,"AuthService")
bullet(doc,"sendOtp(data): POST /gateway/auth/send-otp")
bullet(doc,"verifyOtpAndRegister(data): POST /gateway/auth/verify-register — stores JWT token via TokenService on success")
bullet(doc,"resendOtp(email): POST /gateway/auth/resend-otp")
bullet(doc,"login(data): POST /gateway/auth/login — stores JWT token via TokenService on success")
bullet(doc,"forgotPasswordSendOtp(email): POST /gateway/auth/forgot-password/send-otp")
bullet(doc,"forgotPasswordReset(data): POST /gateway/auth/forgot-password/reset")
bullet(doc,"logout(): Removes token via TokenService; navigates to /auth/login")
bullet(doc,"getProfile(): GET /gateway/auth/profile")
h3(doc,"PolicyService")
bullet(doc,"getPolicyTypes(): GET /gateway/policies/types — public, no auth required")
bullet(doc,"calculatePremium(data): POST /gateway/policies/calculate-premium")
bullet(doc,"createPolicy(data): POST /gateway/policies")
bullet(doc,"getMyPolicies(): GET /gateway/policies/my")
bullet(doc,"getPolicyById(id): GET /gateway/policies/{id}")
bullet(doc,"renewPolicy(policyId, age): POST /gateway/policies/{policyId}/renew")
bullet(doc,"getAdminPolicyTypes(): GET /gateway/policies/admin/types")
bullet(doc,"createPolicyType / updatePolicyType / deletePolicyType / togglePolicyTypeStatus: Admin CRUD operations")
h3(doc,"ClaimService")
bullet(doc,"createClaim(data): POST /gateway/claims")
bullet(doc,"submitClaim(id): POST /gateway/claims/{id}/submit")
bullet(doc,"uploadDocument(id, file): POST /gateway/claims/{id}/documents (multipart/form-data)")
bullet(doc,"deleteDocument(claimId, documentId): DELETE /gateway/claims/{claimId}/documents/{documentId}")
bullet(doc,"getMyClaims(): GET /gateway/claims/my")
bullet(doc,"getClaimById(id): GET /gateway/claims/{id}")
bullet(doc,"updateClaimStatus(claimId, status, adminNote): PUT /gateway/claims/{claimId}/status")
h3(doc,"AdminService")
bullet(doc,"getDashboard(): GET /gateway/admin/dashboard — cached in _dashboardCache")
bullet(doc,"getAllClaims(): GET /gateway/admin/claims — cached in _claimsCache")
bullet(doc,"updateClaimStatus(claimId, status, adminNote): PUT /gateway/admin/claims/status — clears caches")
bullet(doc,"getAllUsers(): GET /gateway/admin/users — cached in _usersCache")
bullet(doc,"updateUserStatus(userId, isActive): PUT /gateway/admin/users/{userId}/status — clears user cache")
bullet(doc,"generateReport(reportType): GET /gateway/admin/reports/generate")
bullet(doc,"getAdminLogs(): GET /gateway/admin/logs")

h2(doc,"8.5 TokenService")
bullet(doc,"Stores JWT in localStorage under key smartsure_token.")
bullet(doc,"Decodes JWT payload using atob() to extract userId, email, fullName, role, and expiry.")
bullet(doc,"isLoggedIn(): Returns true if token exists and is not expired.")
bullet(doc,"isAdmin(): Returns true if decoded role claim equals ADMIN.")
bullet(doc,"getUserId() / getUserName() / getEmail(): Accessors for decoded JWT claims.")

h2(doc,"8.6 Buy Policy Wizard (3-Step Flow)")
numbered(doc,[
    "Step 1 — Details: Customer selects policy type (dropdown shows full coverage details on selection), enters age, start date, and end date. Clicks Calculate Premium.",
    "Step 2 — Premium Breakdown: System displays base amount, age factor amount, duration factor amount, total premium, age factor rules table, and duration factor rules table. Customer clicks Proceed to Confirm.",
    "Step 3 — Confirm and Pay: Summary of policy details and total premium. Customer clicks Pay with Razorpay. Razorpay checkout modal opens. On successful payment, frontend receives razorpay_payment_id and calls POST /gateway/policies to create the policy. Customer is redirected to My Policies.",
])

print("section 8 done")

# SECTION 8 - Key Data Flows
h1(doc,"9. Key Data Flows")

h2(doc,"9.1 User Registration and OTP Verification")
numbered(doc,[
    "User fills Register form. Angular calls POST /gateway/auth/send-otp.",
    "Identity Service: checks email uniqueness (throws EmailAlreadyRegisteredException if exists). Generates 6-digit OTP via RandomNumberGenerator.GetInt32(100000, 1000000). Calls IAuthRepository.UpsertOtpAsync to store with 15-minute expiry. Calls IEmailService.SendOtpEmailAsync via SMTP (Gmail).",
    "Angular navigates to OTP verification screen. User enters OTP code.",
    "Angular calls POST /gateway/auth/verify-register with {email, otpCode, fullName, password}.",
    "Identity Service: validates OTP (checks IsUsed, ExpiresAt, and code match). Creates User with BCrypt-hashed password and role CUSTOMER. Marks OTP as used. JwtHelper generates JWT token.",
    "Response: {token, fullName, email, role, expiresAt}. AuthService stores token via TokenService. User is redirected to Customer Dashboard.",
])

h2(doc,"9.2 Login and JWT Token Flow")
numbered(doc,[
    "Angular calls POST /gateway/auth/login with {email, password}.",
    "Identity Service: retrieves user by email. Verifies BCrypt hash. Checks IsActive. JwtHelper generates JWT with claims: NameIdentifier (userId), Email, Role.",
    "Response: {token, fullName, email, role, expiresAt}. AuthService stores token via TokenService.",
    "AuthInterceptor attaches Bearer token to all subsequent requests automatically.",
    "Route guards read role from TokenService to enforce route-level access control.",
])

h2(doc,"9.3 Policy Purchase Flow")
numbered(doc,[
    "Customer selects policy type. Angular calls GET /gateway/policies/types to populate dropdown with full coverage details.",
    "Customer enters age, start date, end date. Angular calls POST /gateway/policies/calculate-premium. Policy Service returns PremiumResponseDto with full breakdown.",
    "Customer reviews breakdown and proceeds to Step 3.",
    "Customer clicks Pay with Razorpay. Razorpay checkout modal opens in browser.",
    "On successful payment, frontend receives razorpay_payment_id. Angular calls POST /gateway/policies with {policyTypeId, startDate, endDate, age}.",
    "Policy Service: calculates premium, creates Policy with status Active and PolicyNumber=POL-{ticks}, creates Premium record, creates Payment record with TransactionId=TXN-{ticks}.",
    "Customer is redirected to My Policies.",
])

h2(doc,"9.4 Claim Lifecycle Flow")
numbered(doc,[
    "Customer fills claim form. Angular calls POST /gateway/claims. Claims Service creates Claim in Draft status with ClaimNumber=CLM-{ticks}.",
    "Customer uploads documents. Angular calls POST /gateway/claims/{id}/documents (multipart/form-data). FileStorageService saves file to wwwroot/uploads/{claimId}/. ClaimDocument record created.",
    "Customer submits claim. Angular calls POST /gateway/claims/{id}/submit. Claims Service transitions Draft to Submitted. RabbitMQ ClaimSubmitted event published.",
    "Admin views claim in Admin Portal (GET /gateway/admin/claims). Admin opens review panel.",
    "Admin selects new status and enters note. Angular calls PUT /gateway/admin/claims/status.",
    "AdminService pre-fetches claim details. Sends PUT to ClaimsService /api/claims/{id}/status. ClaimsService enforces state machine. Creates AdminLog entry.",
    "AdminService fires background task: fetches customer email from IdentityService. Publishes ClaimStatusNotificationDto to claim.status.notification RabbitMQ queue.",
    "IdentityService ClaimNotificationConsumer receives message. Sends colour-coded HTML email to customer.",
])

h2(doc,"9.5 Policy Renewal Flow")
numbered(doc,[
    "Customer clicks Renew on an Active or Expired policy in My Policies.",
    "Angular calls POST /gateway/policies/{id}/renew with {age}.",
    "Policy Service: validates ownership (PolicyAccessDeniedException if mismatch). Validates status is Active or Expired (PolicyNotRenewableException otherwise). Calculates renewal premium with fixed durationFactor=0.10.",
    "New start date: today if policy is expired, or current EndDate if still active. New end date: +1 year.",
    "Policy updated: new dates, status=Active, IsRenewed=true, RenewalCount+=1. Payment created with TransactionId=TXN-RENEW-{ms}.",
    "Customer is redirected to My Policies with updated policy visible.",
])

# SECTION 9 - Infrastructure
h1(doc,"10. Infrastructure and Configuration")

h2(doc,"10.1 Docker Compose Services")
tbl(doc,
    ["Container","Image","Port(s)","Depends On"],
    [
        ["smartsure-sqlserver","mcr.microsoft.com/mssql/server:2022-latest","1433","—"],
        ["smartsure-rabbitmq","rabbitmq:3-management","5672, 15672","—"],
        ["smartsure-identity","Custom build (IdentityService)","5265","sqlserver (healthy)"],
        ["smartsure-policy","Custom build (PolicyService)","5145","sqlserver (healthy), rabbitmq (healthy)"],
        ["smartsure-claims","Custom build (ClaimsService)","5084","sqlserver (healthy), rabbitmq (healthy)"],
        ["smartsure-admin","Custom build (AdminService)","5073","sqlserver, rabbitmq, identity, policy, claims"],
        ["smartsure-gateway","Custom build (ApiGateway)","5000","identity, policy, claims, admin"],
        ["smartsure-frontend","Custom build (Nginx + Angular)","4200 (→80)","api-gateway"],
    ],
    col_widths=[1.8,2.4,1.2,2.4]
)

h2(doc,"10.2 Database Configuration per Service")
tbl(doc,
    ["Service","Database Name","DbContext Class","Key Tables"],
    [
        ["Identity Service","SmartSure_IdentityDB","AppDbContext","Users, OtpVerifications"],
        ["Policy Service","SmartSure_PolicyDB","PolicyDbContext","PolicyTypes, Policies, Premiums, Payments"],
        ["Claims Service","SmartSure_ClaimsDB","ClaimDbContext","Claims, ClaimDocuments"],
        ["Admin Service","SmartSure_AdminDB","AdminDbContext","AdminLogs, Reports"],
    ],
    col_widths=[1.5,1.8,1.8,2.7]
)

h2(doc,"10.3 JWT Configuration (Shared Across All Services)")
bullet(doc,"JwtSettings:SecretKey — HS256 signing secret. Must be identical across all services.")
bullet(doc,"JwtSettings:Issuer — Token issuer identifier (SmartSure).")
bullet(doc,"JwtSettings:Audience — Token audience identifier (SmartSureClients).")
bullet(doc,"JwtSettings:ExpiryHours — Token lifetime in hours (default 8).")
bullet(doc,"Identity Service generates tokens. All other services validate them using the same parameters.")
bullet(doc,"Token claims: ClaimTypes.NameIdentifier (userId as string), ClaimTypes.Email, ClaimTypes.Role.")

h2(doc,"10.4 RabbitMQ Configuration")
bullet(doc,"All services connect to RabbitMQ on host rabbitmq (Docker) or localhost (local development).")
bullet(doc,"Credentials: smartsure / smartsure123 (Docker); guest / guest (local).")
bullet(doc,"Queue: claim.status.notification — durable, persistent messages.")
bullet(doc,"Publisher: AdminService.Infrastructure.Services.NotificationPublisher — uses RabbitMQ.Client directly.")
bullet(doc,"Consumer: IdentityService.Infrastructure.Messaging.ClaimNotificationConsumer — BackgroundService with auto-reconnect.")

h2(doc,"10.5 Email Configuration (IdentityService)")
bullet(doc,"Provider: Gmail SMTP (smtp.gmail.com:587) with STARTTLS.")
bullet(doc,"Settings: EmailSettings:SmtpHost, SmtpPort, FromEmail, AppPassword, FromName, UseAuthentication, UseStartTls.")
bullet(doc,"Library: MailKit / MimeKit.")
bullet(doc,"Emails sent: OTP for registration, OTP for password reset, claim status change notifications.")

# SECTION 10 - Testing
h1(doc,"11. Unit Test Projects")
tbl(doc,
    ["Test Project","Framework","Tests","Coverage","Key Areas Covered"],
    [
        ["IdentityService.Tests","NUnit 4 + Moq + FluentAssertions","31","~98% line / 100% branch","Registration OTP flow, login, password reset, account deactivation, OTP expiry/reuse"],
        ["PolicyService.Tests","NUnit 4 + Moq + FluentAssertions","34","~97% line / 100% branch","Premium calculation (all age/duration brackets), policy creation, renewal (active/expired/cancelled/wrong user), payment recording"],
        ["ClaimsService.Tests","NUnit 4 + Moq + FluentAssertions","30","~95% line / ~78% branch","Claim creation, submission, status transitions, document upload/delete, access control"],
        ["AdminService.Tests","NUnit 4 + Moq + FluentAssertions","22","~90% line","Dashboard aggregation, claim status update, user status update, report generation, audit logging"],
    ],
    col_widths=[1.8,1.6,0.6,1.2,2.6]
)

h2(doc,"11.1 Testing Approach")
bullet(doc,"Repository interfaces are mocked with MockBehavior.Strict — any unexpected call fails the test immediately.")
bullet(doc,"Service-level mocks use MockBehavior.Loose.")
bullet(doc,"FluentAssertions used for all assertions: .Should().Be(), .Should().ThrowAsync<T>(), .WithMessage().")
bullet(doc,"Custom exception types are asserted directly (e.g. ThrowAsync<PolicyNotFoundException>()) — not the old generic InvalidOperationException.")
bullet(doc,"Premium calculation tests cover all six age/duration combinations from the README examples.")
bullet(doc,"Renewal tests cover: active policy (extends from EndDate), expired policy (starts from today), cancelled policy (throws PolicyNotRenewableException), wrong user (throws PolicyAccessDeniedException), renewal count increment, 1-year duration factor enforcement.")

# SECTION 11 - Security
h1(doc,"12. Security Considerations")
bullet(doc,"Password Storage: Passwords are hashed using BCrypt.Net.BCrypt.HashPassword() and stored in the PasswordHash column of the Users table. The hash cannot be reversed.")
bullet(doc,"JWT Validation: Each service independently validates JWT tokens using the shared signing key. The API Gateway also validates tokens before forwarding to protected routes.")
bullet(doc,"OTP Security: OTPs are 6-digit numeric strings generated via RandomNumberGenerator.GetInt32(100000, 1000000). Expiry is 15 minutes. OtpRecord.IsUsed is set to true after successful verification, preventing reuse.")
bullet(doc,"CORS: Restricted to http://localhost:4200 in the gateway. In Docker, the Angular Nginx container proxies API calls internally.")
bullet(doc,"Authorization: All sensitive endpoints use [Authorize] attribute. Admin-only endpoints use [Authorize(Roles = 'ADMIN')]. Route guards provide an additional client-side layer.")
bullet(doc,"File Upload: Documents are saved to the server filesystem under wwwroot/uploads/{claimId}/. Only the relative URL path is stored in ClaimDocument.FilePath. The wwwroot prefix is stripped from the returned FileUrl.")
bullet(doc,"Global Exception Handling: GlobalExceptionMiddleware in each service catches all unhandled exceptions. Domain exceptions (typed, expected) are logged as Warning. Unexpected exceptions are logged as Error. Stack traces are only included in the response body in the Development environment.")
bullet(doc,"Downstream HTTP Calls: AdminService uses IHttpClientFactory.CreateClient() for background tasks to avoid using a disposed HttpClient after the request scope ends.")

divider(doc)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("SmartSure — Confidential Technical Document  |  May 2026")
rf(r,9,False,(148,163,184),italic=True)

doc.save(r"C:\Users\Pavan\Desktop\SmartSure\SmartSure_LLD.docx")
print("LLD saved successfully")
